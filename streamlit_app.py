import streamlit as st
import pandas as pd
import json
import zipfile
import io
import fnmatch
from datetime import datetime, timedelta
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to format timedelta into hh:mm:ss
def format_timedelta(td):
    days = td.days
    hours, remainder = divmod(td.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

# Function to parse and format text with <B> and <I> tags
def parse_and_add_run(paragraph, text):
    while text:
        if text.startswith("<B>"):
            text = text[3:]  # Remove the <B> tag
            end_index = text.find("</B>")
            if end_index != -1:
                paragraph.add_run(text[:end_index]).bold = True
                text = text[end_index + 4:]  # Remove the </B> tag
            else:
                paragraph.add_run(text).bold = True
                break
        elif text.startswith("<I>"):
            text = text[3:]  # Remove the <I> tag
            end_index = text.find("</I>")
            if end_index != -1:
                paragraph.add_run(text[:end_index]).italic = True
                text = text[end_index + 4:]  # Remove the </I> tag
            else:
                paragraph.add_run(text).italic = True
                break
        else:
            next_tag = min(
                (text.find("<B>"), text.find("<I>")),
                key=lambda x: (x == -1, x)  # Find the next tag or end
            )
            if next_tag == -1:
                paragraph.add_run(text)
                break
            paragraph.add_run(text[:next_tag])
            text = text[next_tag:]

# Function to create a cumulative timing table
def create_cumulative_timing_table(facilitator_content, start_time):
    cumulative_elapsed_time = timedelta(0)  # Start with zero elapsed time
    table_data = []

    for i, item in enumerate(facilitator_content):
        cumulative_elapsed_time += timedelta(seconds=item.get('timer_seconds', 0))
        cumulative_time = start_time + cumulative_elapsed_time
        formatted_time = format_timedelta(cumulative_elapsed_time)  # Format as hh:mm:ss
        
        table_data.append({
            'Cumulative Time (hh:mm:ss)': formatted_time,
            'Subject': item.get('subject', ''),
            'Text': item.get('text', ''),
            'Inject Timing (s)': item.get('timer_seconds', 0)
        })

    return pd.DataFrame(table_data)

# Function to add shading to a cell
def add_shading(cell, color):
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tc_pr.append(shd)

# Function to save the cumulative timing table to a Word document
def save_to_word(table, mel_name):
    doc = Document()
    doc.add_heading(f'Cumulative Timing Table for {mel_name}', 0)

    table_to_word = doc.add_table(rows=1, cols=4)
    hdr_cells = table_to_word.rows[0].cells
    hdr_cells[0].text = 'Cumulative Time (hh:mm:ss)'
    hdr_cells[1].text = 'Subject'
    hdr_cells[2].text = 'Text'
    hdr_cells[3].text = 'Inject Timing (s)'

    # Add shading to header
    for cell in hdr_cells:
        add_shading(cell, 'D3D3D3')  # Light grey for header

    # Add rows with alternating shading
    for i, row in table.iterrows():
        row_cells = table_to_word.add_row().cells
        row_cells[0].text = row['Cumulative Time (hh:mm:ss)']
        parse_and_add_run(row_cells[1].paragraphs[0], row['Subject'])
        parse_and_add_run(row_cells[2].paragraphs[0], row['Text'])
        row_cells[3].text = str(row['Inject Timing (s)'])

        if i % 2 == 0:
            for cell in row_cells:
                add_shading(cell, 'D3D3D3')  # Light grey

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to retrieve facilitator content for each MEL
def get_facilitator_content_for_mel(stages, mel_id):
    facilitator_content = []

    for stage in stages:
        # Debugging: Print each stage's mel_id and the target mel_id
        st.write(f"Checking stage {stage.get('id')} with mel_id: {stage.get('mel_id')} against target mel_id: {mel_id}")
        
        if str(stage.get('mel_id')) == str(mel_id):  # Ensure both are compared as strings
            # Determine the inject timing
            timer_seconds = 0
            if stage.get('timer_answers'):
                for timer in stage['timer_answers']:
                    if 'timer_seconds' in timer:
                        timer_seconds = timer['timer_seconds']
                        break

            stage_info = {
                'subject': stage.get('subject'),
                'text': stage.get('text'),
                'timer_seconds': timer_seconds
            }
            facilitator_content.append(stage_info)

    # Debugging: Check if any stages were added
    st.write(f"Facilitator content for MEL ID {mel_id}: {len(facilitator_content)} stages found")

    return facilitator_content


# Streamlit app
def main():
    st.title("Cumulative Timing Tables from .txplib File")

    # Ask the user for a start time
    start_time_input = st.text_input("Enter the start time (format: HH:MM:SS)", value="00:00:00")
    start_time = datetime.strptime(f"2024-01-01 {start_time_input}", "%Y-%m-%d %H:%M:%S")

    # Upload the .txplib file
    uploaded_file = st.file_uploader("Choose a .txplib file", type="txplib")
    
    if uploaded_file is not None:
        st.write("File uploaded successfully.")
        
        try:
            with zipfile.ZipFile(io.BytesIO(uploaded_file.read()), 'r') as zip_ref:
                st.write("Zip file opened successfully.")
                
                design_files = []
                for file_name in zip_ref.namelist():
                    if fnmatch.fnmatch(file_name, "design *.txt"):
                        st.write(f"Found matching design file: {file_name}")
                        with zip_ref.open(file_name) as json_file:
                            json_file_content = json_file.read()
                            design_files.append((file_name, json_file_content))

                if design_files:
                    for file_name, json_file_content in design_files:
                        data = json.loads(json_file_content)
                        stages = data.get('stages', [])
                        mels = data.get('mels', [])

                        for mel in mels:
                            mel_id = mel['id']
                            mel_name = mel['name']
                            st.write(f"Processing MEL: {mel_name}")

                            facilitator_content = get_facilitator_content_for_mel(stages, mel_id)
                            cumulative_timing_table = create_cumulative_timing_table(facilitator_content, start_time)

                            st.write(f"Displaying the cumulative timing table for {mel_name}")
                            st.dataframe(cumulative_timing_table)

                            if st.button(f"Download {mel_name} table as Word"):
                                word_buffer = save_to_word(cumulative_timing_table, mel_name)
                                st.download_button(
                                    label=f"Download {mel_name} table as Word",
                                    data=word_buffer,
                                    file_name=f"{mel_name}_timing_table.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                else:
                    st.error("No valid 'design *.txt' files found inside the .txplib archive.")
        except Exception as e:
            st.error(f"Error processing the zip file: {e}")

if __name__ == "__main__":
    main()
